# import streamlit as st
# import pdfplumber
# import docx
# import re
# import numpy as np
# from sklearn.feature_extraction.text import TfidfVectorizer
# from sentence_transformers import SentenceTransformer
# from sklearn.metrics.pairwise import cosine_similarity
# import matplotlib.pyplot as plt
# from io import BytesIO
#
# # ------------------------------
# # Utility Functions
# # ------------------------------
#
# def extract_text_from_pdf(file):
#     text = ""
#     with pdfplumber.open(file) as pdf:
#         for page in pdf.pages:
#             text += page.extract_text() + "\n"
#     return text
#
#
# def extract_text_from_docx(file):
#     doc = docx.Document(file)
#     return "\n".join([para.text for para in doc.paragraphs])
#
#
# def clean_text(text):
#     text = text.lower()
#     text = re.sub(r'[^a-zA-Z0-9\s]', '', text)
#     return text
#
#
# def extract_keywords(text, top_n=30):
#     vectorizer = TfidfVectorizer(stop_words='english', max_features=top_n)
#     X = vectorizer.fit_transform([text])
#     return vectorizer.get_feature_names_out()
#
#
# def keyword_match_score(resume_text, jd_keywords):
#     matched = []
#     missing = []
#     for word in jd_keywords:
#         if word in resume_text:
#             matched.append(word)
#         else:
#             missing.append(word)
#     score = (len(matched) / len(jd_keywords)) * 100 if jd_keywords.any() else 0
#     return score, matched, missing
#
#
# def semantic_similarity_score(resume_text, jd_text):
#     model = SentenceTransformer('all-MiniLM-L6-v2')
#     embeddings = model.encode([resume_text, jd_text])
#     score = cosine_similarity([embeddings[0]], [embeddings[1]])[0][0]
#     return score * 100
#
#
# def calculate_final_score(keyword_score, semantic_score):
#     return (0.4 * keyword_score) + (0.6 * semantic_score)
#
#
# def show_gauge(score):
#     fig, ax = plt.subplots()
#     ax.barh(0, score)
#     ax.set_xlim(0, 100)
#     ax.set_title(f"ATS Score: {round(score,2)}%")
#     ax.set_yticks([])
#     st.pyplot(fig)
#
#
# # ------------------------------
# # Streamlit UI
# # ------------------------------
#
# st.set_page_config(page_title="ATS Resume Analyzer", layout="wide")
# st.title("🚀 AI ATS Resume Analyzer")
#
# st.write("Upload your Resume and paste Job Description to evaluate ATS compatibility.")
#
# resume_file = st.file_uploader("Upload Resume (PDF or DOCX)", type=["pdf", "docx"])
# job_description = st.text_area("Paste Job Description Here")
#
# if st.button("Analyze"):
#
#     if resume_file and job_description:
#
#         # Extract resume text
#         if resume_file.type == "application/pdf":
#             resume_text = extract_text_from_pdf(resume_file)
#         else:
#             resume_text = extract_text_from_docx(resume_file)
#
#         resume_text_clean = clean_text(resume_text)
#         jd_text_clean = clean_text(job_description)
#
#         # Extract JD Keywords
#         jd_keywords = extract_keywords(jd_text_clean)
#
#         # Keyword Score
#         keyword_score, matched, missing = keyword_match_score(resume_text_clean, jd_keywords)
#
#         # Semantic Score
#         semantic_score = semantic_similarity_score(resume_text_clean, jd_text_clean)
#
#         # Final ATS Score
#         final_score = calculate_final_score(keyword_score, semantic_score)
#
#         # Classification
#         if final_score >= 85:
#             status = "🟢 Highly ATS Friendly"
#         elif final_score >= 65:
#             status = "🟡 Moderately ATS Friendly"
#         else:
#             status = "🔴 Needs Optimization"
#
#         # Dashboard Layout
#         col1, col2 = st.columns(2)
#
#         with col1:
#             st.subheader("ATS Score Compass")
#             show_gauge(final_score)
#
#         with col2:
#             st.subheader("Score Breakdown")
#             st.write(f"Keyword Match Score: {round(keyword_score,2)}%")
#             st.write(f"Semantic Similarity Score: {round(semantic_score,2)}%")
#             st.write(f"Final ATS Score: {round(final_score,2)}%")
#             st.write(status)
#
#         st.subheader("Matched Keywords")
#         st.write(", ".join(matched) if matched else "None")
#
#         st.subheader("Missing Keywords")
#         st.write(", ".join(missing) if missing else "None 🎉")
#
#     else:
#         st.warning("Please upload resume and paste job description.")
# ==============================
# Advanced ATS Resume Analyzer
# Designed & Developed by Kajal Dadas
# ==============================

import streamlit as st
import pdfplumber
import docx
import re
import numpy as np
import matplotlib.pyplot as plt
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer

# ------------------------------
# CONFIG
# ------------------------------

st.set_page_config(page_title="ATS Resume Intelligence", layout="wide")

# ------------------------------
# TEXT EXTRACTION
# ------------------------------

def extract_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            if page.extract_text():
                text += page.extract_text()
    return text


def extract_docx(file):
    doc = docx.Document(file)
    return "\n".join([p.text for p in doc.paragraphs])


def clean_text(text):
    text = text.lower()
    text = re.sub(r'[^a-z0-9\s]', '', text)
    return text


# ------------------------------
# KEYWORD EXTRACTION
# ------------------------------

def extract_keywords(text, top_n=40):
    vectorizer = TfidfVectorizer(stop_words='english', max_features=top_n)
    vectorizer.fit([text])
    return vectorizer.get_feature_names_out()


def keyword_score(resume, jd_keywords):
    matched = []
    missing = []
    for word in jd_keywords:
        if word in resume:
            matched.append(word)
        else:
            missing.append(word)

    score = (len(matched) / len(jd_keywords)) * 100 if len(jd_keywords) > 0 else 0
    return score, matched, missing


# ------------------------------
# SEMANTIC SCORE
# ------------------------------

def semantic_score(resume, jd):
    model = SentenceTransformer("all-MiniLM-L6-v2")
    emb = model.encode([resume, jd])
    score = cosine_similarity([emb[0]], [emb[1]])[0][0]
    return score * 100


# ------------------------------
# VISUAL GAUGE
# ------------------------------

def draw_gauge(score):
    fig, ax = plt.subplots(figsize=(6, 3))

    ax.barh(0, score, height=0.5)
    ax.set_xlim(0, 100)
    ax.set_yticks([])
    ax.set_title("ATS Compatibility Score", fontsize=14, fontweight="bold")

    ax.text(score / 2, 0, f"{round(score,2)}%",
            va='center', ha='center', fontsize=18, fontweight='bold')

    st.pyplot(fig)


# ------------------------------
# RADAR CHART
# ------------------------------

def radar_chart(keyword_s, semantic_s):
    categories = ['Keyword Match', 'Semantic Similarity']
    values = [keyword_s, semantic_s]

    angles = np.linspace(0, 2 * np.pi, len(categories), endpoint=False)
    values = np.concatenate((values, [values[0]]))
    angles = np.concatenate((angles, [angles[0]]))

    fig = plt.figure(figsize=(4,4))
    ax = fig.add_subplot(111, polar=True)
    ax.plot(angles, values)
    ax.fill(angles, values, alpha=0.25)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories)
    ax.set_yticklabels([])

    st.pyplot(fig)


# ------------------------------
# UI DESIGN
# ------------------------------

st.markdown("""
    <h1 style='text-align: center;'>🚀 ATS Resume Intelligence Dashboard</h1>
    <p style='text-align: center; font-size:14px;'>Designed & Developed by <b>Kajal Dadas</b></p>
    <hr>
""", unsafe_allow_html=True)

col1, col2 = st.columns(2)

with col1:
    resume_file = st.file_uploader("Upload Resume (PDF/DOCX)", type=["pdf", "docx"])

with col2:
    jd_text = st.text_area("Paste Job Description")

if st.button("Analyze Resume"):

    if resume_file and jd_text:

        # Extract Resume Text
        if resume_file.type == "application/pdf":
            resume_text = extract_pdf(resume_file)
        else:
            resume_text = extract_docx(resume_file)

        resume_clean = clean_text(resume_text)
        jd_clean = clean_text(jd_text)

        # Keyword Analysis
        jd_keywords = extract_keywords(jd_clean)
        keyword_s, matched, missing = keyword_score(resume_clean, jd_keywords)

        # Semantic Analysis
        semantic_s = semantic_score(resume_clean, jd_clean)

        # Final Score
        final_score = (0.4 * keyword_s) + (0.6 * semantic_s)

        # Status
        if final_score >= 85:
            status = "🟢 Highly ATS Friendly"
        elif final_score >= 65:
            status = "🟡 Moderately ATS Friendly"
        else:
            status = "🔴 Needs Optimization"

        st.markdown("---")

        # Dashboard Layout
        c1, c2 = st.columns(2)

        with c1:
            draw_gauge(final_score)
            st.markdown(f"<h3 style='text-align:center;'>{status}</h3>", unsafe_allow_html=True)

        with c2:
            radar_chart(keyword_s, semantic_s)

        st.markdown("### 📊 Score Breakdown")
        st.metric("Keyword Match %", f"{round(keyword_s,2)}%")
        st.metric("Semantic Similarity %", f"{round(semantic_s,2)}%")
        st.metric("Final ATS Score %", f"{round(final_score,2)}%")

        st.markdown("---")

        st.markdown("### ✅ Matched Keywords")
        st.write(", ".join(matched) if matched else "None")

        st.markdown("### ❌ Missing Keywords")
        st.write(", ".join(missing) if missing else "None 🎉")

    else:
        st.warning("Please upload resume and paste job description.")
